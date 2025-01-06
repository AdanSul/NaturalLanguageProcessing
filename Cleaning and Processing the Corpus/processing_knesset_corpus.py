# imports
from docx import Document
import sys
import os
import re
import json

# Global list of exception words
EXCEPTION_WORDS = [
   'משתתפים', 'מנהלות הוועדה', 'מנהל הוועדה', 'מוזמנים באמצעים מקוונים','משרד האוצר',
   'משתתפים באמצעים מקוונים', 'נכחו', 'מנהלת הוועדה', 'מסקנות הוועדה', 'משתתפים',
   'מוזמנים', 'סדר היום', 'חברי הוועדה', 'חברי הכנסת', 'ייעוץ משפטי',
   'מנהל/ת הוועדה', 'רצוני לשאול', 'קריאות', 'קריאה', 'האפשרות', 'ביום',
   'ברצוני לשאול', 'רישום פרלמנטרי','המשרד מתקצב חמש תכניות', 'נוכחים','משרד המשפטים','יועץ משפטי',
   'סדר-היום','רשמת פרלמנטרית', 'קצרנית','מנהלי הוועדה','חברי כנסת','רשמה וערכה'
]

DIFFERENT_FILES=[
    '25_ptv_1219728.docx','23_ptv_600338.docx','25_ptv_1457545.docx','25_ptv_3841247.docx','23_ptv_599659.docx',
    '23_ptv_598723.docx','23_ptv_598323.docx','23_ptv_582824.docx','20_ptv_519812.docx',
    '20_ptv_490139.docx','20_ptv_488037.docx','20_ptv_397418.docx','20_ptv_387379.docx','20_ptv_370910.docx',
    '20_ptv_341230.docx','20_ptv_341020.docx','20_ptv_320584.docx','20_ptv_311936.docx','19_ptv_266962.docx',
    '19_ptv_232326.docx','19_ptv_302840.docx','19_ptv_279660.docx','19_ptv_262672.docx']

def extract_metadata(file_name):
    # Remove file extension and split by underscores
    base_name = os.path.splitext(file_name)[0]
    parts = base_name.split("_")

    # Extract knesset number, protocol type, and protocol number
    knesset_number = int(parts[0])
    protocol_type = parts[1]

    # Map protocol type to a meaningful name
    protocol_type_mapping = {
        "ptm": "plenary",
        "ptv": "committee"
    }

    protocol_type = protocol_type_mapping.get(protocol_type, "unknown")
    return knesset_number, protocol_type

hebrew_to_int = {
    'החמש עשרה': 15, 'החמישים וארבע': 54, 'השלושים ושתיים': 32, 'הארבעים ושש': 46,
    'הארבעים ותשע': 49, 'השישים ותשע': 69, 'השישים ושלוש': 63, 'החמישים ושבע': 57, 'המאה ועשרים וארבע': 124, 
    'המאה ושמונים וארבע': 184, 'המאה וחמישים ושש': 156, 'המאתיים ושמונים ושתיים': 282, 'המאתיים ושישים ותשע': 269,
    'המאה ושמונים וחמש': 185, 'המאתיים ושמונים': 280, 'השלוש מאות ושתיים': 302, 'השלוש מאות ושישים ושש': 366,
    'השלוש מאות ושמונים וארבע': 384, 'המאתיים ושישים ושבע': 267, 'המאה ושמונים ושלוש': 183,
    'הארבעים': 40, 'השמונים ותשע': 89, 'המאה וארבעים ואחת': 141, 'המאתיים ושישים ושתיים': 262, 'התשעים ושתיים': 92
}

def extract_protocol_number(text):
    # Match numerical protocol numbers (e.g., "פרוטוקול מס' 45")
    protocol_number_match = re.search(r'פרוטוקול מס\'\s*(\d+(?:,\s*\d+)*)', text)
    if protocol_number_match:
        numbers = protocol_number_match.group(1).split(",")  # Split numbers by comma
        for number in numbers:
            try:
                return int(number.strip())  # Return the first valid number
            except ValueError:
                pass

    # Match Hebrew text-based session numbers (e.g., "הישיבה המאה וחמישים ושש של הכנסת")
    hebrew_session_match = re.search(r'הישיבה\s+(.*?)\s+של\s+הכנסת', text)
    if hebrew_session_match:
        hebrew_number_text = hebrew_session_match.group(1).replace("-", " ")
        if hebrew_number_text in hebrew_to_int:
            return hebrew_to_int[hebrew_number_text]

    # If no valid number is found, return -1
    return -1

def clean_speaker_name(raw_name):
    # Remove text within parentheses
    cleaned_name = re.sub(r'\([^)]*\)', '', raw_name)

    if cleaned_name.startswith('ביום') or cleaned_name.startswith('אני') or cleaned_name.startswith('בפסקה') or cleaned_name.startswith('מר'):
        return None
    
    if cleaned_name.startswith('משתתפים') or cleaned_name.startswith('האפשרות') or cleaned_name.startswith('קוראת') or cleaned_name.startswith('כרצוני לשאול') :
        return None 

    # Remove certain patterns (titles, prefixes, and suffixes)
    patterns_to_remove = [
        r'\s*סגנית מזכיר הכנסת\s+', r'\s*סגן מזכיר הכנסת\s+', r'\s*סגן מזכירת הכנסת\s+',
        r'\s*סגנית מזכירת הכנסת\s+', r'\s*מזכיר הכנסת\s+', r'\s*מזכירת הכנסת\s+',
        r'\s*סגנית שר במשרד ראש הממשלה\s+', r'\s*סגן שר במשרד ראש הממשלה\s+',
        r'\s*השר לנושאים אסטרטגיים ולענייני מודיעין\s+', r'\s*(סגן|סגנית)\s+(ה\w+\s+וה\w+)\s*', 
        r'\s*(סגן|סגנית)\s+(ה\w+\s+)\s*', r'\s*סגן\s+', r'\s*סגנית\s+', r'(?:שרת?\s+ה\w+,\s+ה\w+\s+וה\w+\s+).*',
        r'\s*שר(?:ת)?\s+ה\w+\s+וה\w+\s*', r'\s*שר(?:ת)?\s+ה\w+\s*', r'\s*השר(?:ה)?\s+ל\w+\s+\w+\s+', 
        r'\s*שר(?:ת)?\s+ל\w+\s+[^,]\s', r'\s*סגן שר(?:ת)?(?:\s+[^,])?,', r'\s*שר(?:ת)?(?:\s+[^,])?,'         
    ]

    for pattern in patterns_to_remove:
        cleaned_name = re.sub(pattern, '', cleaned_name)

    # Remove specific predefined titles and phrases
    phrases_to_remove = [
        '<< יור >>', '<< דובר >>', '(יש עתיד-תל"ם)', '<< אורח >>',
        '<< דובר_המשך >>', '<< קריאה >>', '<', '>', 'מ"מ היו"ר', 'היו"ר','היו”ר','היו”ר '
        'השר', 'לביטחון פנים', 'ראש הממשלה', 'הלאומיות', 'פרופ',
        ' מר ', 'ופיתוח הכפר', 'תשובת','יו"ר', ' לאיכות הסביבה', 'עו"ד', 'נצ"מ', 'היו"ר '
        'ד"ר'
    ]
    for phrase in phrases_to_remove:
        cleaned_name = cleaned_name.replace(phrase, "")

    if cleaned_name in EXCEPTION_WORDS:
        return None
    
    # Final cleanup: strip extra spaces and return
    return cleaned_name.strip()

def extract_speaker_name(text, file_name):
    global EXCEPTION_WORDS

    if file_name not in DIFFERENT_FILES:
        # Regex to capture the speaker's name before the colon
        speaker_pattern = r'^([^\n:]+?):'
        speaker_match = re.search(speaker_pattern, text)

        # Check if there are words following the colon
        words_after_colon_pattern = r':\s*(\w+)'
        words_after_colon_match = re.search(words_after_colon_pattern, text)

        # Check if the colon is followed by the end of the line or end of the string
        colon_at_end_pattern = r':\s*$'
        colon_at_end_match = re.search(colon_at_end_pattern, text)

        # If a valid speaker name is found
        if speaker_match and colon_at_end_match and not words_after_colon_match:
            raw_speaker_name = speaker_match.group(1).strip()

            # Handle normal cases
            if raw_speaker_name in EXCEPTION_WORDS  or re.match(r'^סדר[-\s]היום', raw_speaker_name):
                return None  # Skip if it's in the exception list

            return clean_speaker_name(raw_speaker_name)

        return None
    
    else:
        speaker_pattern =  r'^([^\n:]+?):'
        speaker_match = re.search(speaker_pattern, text)

        if  '<' in text and '>' in text and ':' in text and not any(char.isdigit() for char in text) and speaker_match is not None:
            raw_speaker_name = speaker_match.group(1).strip()
            if raw_speaker_name in EXCEPTION_WORDS:
                return None
            cleaned_name = re.sub(r'\([^)]*\)','', raw_speaker_name)

            phrases_to_remove = ['<< יור >>', "<< דובר >>", '(יש עתיד-תל"ם)', '<< אורח >>',
            '<< דובר_המשך >>', '<< קריאה >>', '<', '>','היו"ר' ,'שר הבינוי והשיכון '    
            ]

            for phrase in phrases_to_remove:
                cleaned_name = cleaned_name.replace(phrase, "")
            
            cleaned_name = cleaned_name.strip()

            if cleaned_name in EXCEPTION_WORDS:
                return None
            
            return cleaned_name

def save_data_to_json(data_list, output_path):
    with open(output_path, "w", encoding="utf-8") as json_file:
        for data in data_list:
            json_file.write(json.dumps(data, ensure_ascii=False) + "\n")

def break_text_into_phrases(input_text):
    buffer = ''
    separators = {'.', '!', '?'}
    phrases = []

    for i, symbol in enumerate(input_text):
        if symbol in separators:
            # Check if the separator is part of a number or date
            if i > 0 and i < len(input_text) - 1:
                prev_char = input_text[i - 1]
                next_char = input_text[i + 1]

                if prev_char.isdigit() and next_char.isdigit():
                    buffer += symbol  # Treat as part of a number/date
                    continue
            
            # Check if the period is part of a numbered list
            if i > 0 and input_text[i - 1].isdigit() and input_text[i - 2] == ' ' and (i < len(input_text) - 1 and input_text[i + 1] == ' '):
                buffer += symbol  # Treat as part of the enumeration
                continue

            # Add the sentence to the list
            buffer += symbol
            if buffer.strip():
                phrases.append(buffer.strip())
            buffer = ''
        else:
            buffer += symbol

    if buffer.strip():
        phrases.append(buffer.strip())
    
    return phrases

def clean_sentence(sentence):
    words = re.findall(r'\b[א-ת]+\b', sentence)  # מוצא רק מילים בעברית
    if len(words) <= 2:
        return None

    if '- - -' in sentence:
        return None
    if '- -' in sentence:
        return None
    if '- - - - - - - - -' in sentence:
        return None
    if '---' in sentence:
        return None
    if '--' in sentence:
        return None
    if '– – –' in sentence:
        return None
    if '- - - - - - - - -' in sentence:
        return None

    # הסרת מילים באנגלית
    cleaned_sentence = re.sub(r'\b[a-zA-Z]+\b', '', sentence)

    # הסרת רווחים מיותרים
    cleaned_sentence = re.sub(r'\s+', ' ', cleaned_sentence).strip()

    # הסרת שורות ריקות ושחזור הטקסט
    cleaned_sentence = "\n".join([line for line in cleaned_sentence.split('\n') if line.strip()])

    # בדיקת משפטים המכילים אותיות עבריות בלבד
    if not re.search(r'[א-ת]', cleaned_sentence):
        return None

    return cleaned_sentence



def get_tokens(text):
    tokens = []
    words = text.split(" ")
    for word in words:
        tokens.append(word)
    return tokens

def tokenize_sentence(sentence):
    split_punctuations = '.,;!?()[]{}-"\''  # סימני פיסוק להפרדה
    tokenized_sentence = []
    word = ""
    hebrew_abbreviation = False

    for i, char in enumerate(sentence):
        # אם מדובר בראשי תיבות כמו "ד"ר", השאר שלם
        if char == "״" or (char == '"' and i > 0 and sentence[i - 1].isalpha()):
            hebrew_abbreviation = True
            word += char
            continue
        elif hebrew_abbreviation and char.isalpha():
            word += char
            hebrew_abbreviation = False
            continue
        # אם מדובר בסימן פיסוק
        if char in split_punctuations:
            if word:
                tokenized_sentence.append(word)
                word = ""
            tokenized_sentence.append(char)
        elif char.isspace():
            if word:
                tokenized_sentence.append(word)
                word = ""
        else:
            word += char
    if word:
        tokenized_sentence.append(word)
    return " ".join(tokenized_sentence)


def process_documents(folder_path, output_file):
    # Get all .docx files in the folder
    files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

    data_list = []

    for file_name in files:
        knesset_number, protocol_type = extract_metadata(file_name)
        doc_path = os.path.join(folder_path, file_name)

        try:
            doc = Document(doc_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            protocol_number = extract_protocol_number(text)

            current_speaker = None
            current_text = []

            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if not paragraph_text:
                    continue

                speaker_name = extract_speaker_name(paragraph_text, file_name)

                if speaker_name and len(speaker_name.split()) <= 4:
                    if current_speaker and current_text:
                        full_text = " ".join(current_text)
                        sentences = break_text_into_phrases(full_text)

                        for sentence in sentences:
                            clean_sentence_text = clean_sentence(sentence)
                            if clean_sentence_text:
                                tokenized_sentence = tokenize_sentence(clean_sentence_text)
                                tokens = get_tokens(clean_sentence_text)  # Tokenize the cleaned sentence
                                if len(tokens) >= 4:  # Only add if there are 4 or more tokens
                                    data_list.append({
                                        "protocol_name": file_name,
                                        "knesset_number": knesset_number,
                                        "protocol_type": protocol_type,
                                        "protocol_number": protocol_number,
                                        "speaker_name": current_speaker,
                                        "sentence_text": tokenized_sentence
                                    })

                    current_speaker = speaker_name
                    current_text = [paragraph_text.split(":", 1)[1].strip()] if ":" in paragraph_text else []
                else:
                    if current_speaker:
                        current_text.append(paragraph_text)

            if current_speaker and current_text:
                full_text = " ".join(current_text)
                sentences = break_text_into_phrases(full_text)

                for sentence in sentences:
                    clean_sentence_text = clean_sentence(sentence)
                    if clean_sentence_text:
                        tokenized_sentence = tokenize_sentence(clean_sentence_text)
                        tokens = get_tokens(clean_sentence_text)  # Tokenize the cleaned sentence
                        if len(tokens) >= 4:  # Only add if there are 4 or more tokens
                            data_list.append({
                                "protocol_name": file_name,
                                "knesset_number": knesset_number,
                                "protocol_type": protocol_type,
                                "protocol_number": protocol_number,
                                "speaker_name": current_speaker,
                                "sentence_text": tokenized_sentence
                            })

        except Exception:
            pass

    save_data_to_json(data_list, output_file)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit(1)

    # Get input folder and output file from command-line arguments
    input_folder = sys.argv[1]
    output_file = sys.argv[2]

    # Run the main function
    process_documents(input_folder, output_file)


